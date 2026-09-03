"""
Google DriveのファイルをPineconeに登録するローダー
"""
import sys
import os
import io

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../..')))

from langchain_core.documents import Document
from googleapiclient.http import MediaIoBaseDownload
from src.auth.google_auth import get_google_drive_service
from src.loaders.pinecone_storage import save_to_pinecone as _save_to_pinecone
from config.settings import settings
import logging

# PDF/Docx処理用
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class GoogleDriveLoader:
    """
    Google DriveのファイルをPineconeに登録するクラス
    """

    # サポートするMIMEタイプ
    SUPPORTED_MIME_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.google-apps.document': 'google_doc',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'text/plain': 'text',
        'text/markdown': 'markdown',
    }

    def __init__(self, credentials_path: str = None, token_path: str = None):
        """
        Args:
            credentials_path: OAuth 2.0 クライアント認証情報JSONファイルのパス
            token_path: 認証トークンを保存するパス
        """
        self.credentials_path = credentials_path or settings.google_drive_credentials_path
        self.token_path = token_path or settings.google_drive_token_path

        # Google Drive APIサービスを初期化
        self.service = get_google_drive_service(self.credentials_path, self.token_path)

    def list_files_in_folder(self, folder_id: str = None, recursive: bool = True) -> list:
        """
        指定されたフォルダ内のファイル一覧を取得

        Args:
            folder_id: Google DriveのフォルダID（Noneの場合は全ファイル）
            recursive: サブフォルダも再帰的に検索するか

        Returns:
            list: ファイル情報のリスト
        """
        all_files = []

        try:
            # フォルダ内のアイテムを取得（フォルダも含む）
            query_parts = []
            if folder_id:
                query_parts.append(f"'{folder_id}' in parents")
            query_parts.append("trashed=false")
            query = " and ".join(query_parts)

            results = self.service.files().list(
                q=query,
                pageSize=100,
                fields="files(id, name, mimeType, modifiedTime, webViewLink)",
                supportsAllDrives=True,
                includeItemsFromAllDrives=True,
                corpora='allDrives'
            ).execute()

            items = results.get('files', [])

            for item in items:
                mime_type = item['mimeType']

                # フォルダの場合は再帰的に検索
                if mime_type == 'application/vnd.google-apps.folder':
                    if recursive:
                        logger.info(f"Entering folder: {item['name']}")
                        sub_files = self.list_files_in_folder(item['id'], recursive=True)
                        all_files.extend(sub_files)
                # サポートするファイルタイプの場合はリストに追加
                elif mime_type in self.SUPPORTED_MIME_TYPES:
                    all_files.append(item)

            logger.info(f"Found {len(all_files)} files in Google Drive")
            return all_files

        except Exception as e:
            logger.error(f"Failed to list files: {e}")
            return []

    def download_file_content(self, file_id: str, mime_type: str) -> str:
        """
        Google Driveからファイルの内容をダウンロード
        """
        try:
            # PDFはGoogle Drive OCR経由でテキスト抽出、失敗時はpypdfでフォールバック
            if mime_type == 'application/pdf':
                text = self._extract_pdf_text_via_ocr(file_id)
                if text is not None:
                    return text
                # OCR失敗 → pypdfフォールバック
                logger.info(f"Falling back to pypdf for {file_id}")
                try:
                    request = self.service.files().get_media(fileId=file_id, supportsAllDrives=True)
                    buf = io.BytesIO()
                    downloader = MediaIoBaseDownload(buf, request)
                    done = False
                    while not done:
                        _, done = downloader.next_chunk()
                    buf.seek(0)
                    reader = PdfReader(buf)
                    return "\n".join(page.extract_text() or "" for page in reader.pages)
                except Exception as e:
                    logger.error(f"pypdf fallback also failed for {file_id}: {e}")
                    return ""

            # Google Docsの場合はエクスポート
            if mime_type == 'application/vnd.google-apps.document':
                request = self.service.files().export_media(
                    fileId=file_id,
                    mimeType='text/plain'
                )
            else:
                request = self.service.files().get_media(
                    fileId=file_id,
                    supportsAllDrives=True
                )

            file_buffer = io.BytesIO()
            downloader = MediaIoBaseDownload(file_buffer, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()
            file_buffer.seek(0)

            if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return self._extract_docx_text(file_buffer)
            else:
                return file_buffer.read().decode('utf-8')

        except Exception as e:
            logger.error(f"Failed to download file {file_id}: {e}")
            return ""

    def _extract_pdf_text_via_ocr(self, file_id: str) -> str:
        """Google Drive OCRを使ってPDFからテキストを抽出"""
        doc_id = None
        try:
            # PDFをGoogle Docsとしてコピー（OCR変換）
            # parentsにrootを指定してMy Driveにコピー（共有ドライブだとdrive.fileスコープで削除できないため）
            copy_response = self.service.files().copy(
                fileId=file_id,
                body={'mimeType': 'application/vnd.google-apps.document', 'parents': ['root']},
                supportsAllDrives=True
            ).execute()
            doc_id = copy_response['id']

            # Google DocsをプレーンテキストとしてExport
            request = self.service.files().export_media(
                fileId=doc_id,
                mimeType='text/plain'
            )
            file_buffer = io.BytesIO()
            downloader = MediaIoBaseDownload(file_buffer, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()
            file_buffer.seek(0)
            text = file_buffer.read().decode('utf-8')
            logger.info(f"OCR extraction succeeded: {len(text)} chars")
            return text

        except Exception as e:
            logger.error(f"OCR extraction failed for {file_id}: {e}")
            return None  # Noneでフォールバックを示す
        finally:
            # 一時的に作成したGoogle Docsを削除
            if doc_id:
                try:
                    self.service.files().update(fileId=doc_id, body={'trashed': True}).execute()
                    logger.info(f"Trashed temp OCR doc: {doc_id}")
                except Exception as e:
                    logger.warning(f"Failed to trash temp OCR doc {doc_id}: {e}. Please delete it manually from Google Drive.")

    def _extract_docx_text(self, file_buffer: io.BytesIO) -> str:
        """Docxからテキストを抽出"""
        try:
            doc = DocxDocument(file_buffer)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {e}")
            return ""

    def load_file(self, file_info: dict) -> Document:
        """
        単一ファイルを読み込み

        Args:
            file_info: ファイル情報（list_files_in_folderの戻り値の要素）

        Returns:
            Document: LangChainのDocumentオブジェクト
        """
        try:
            file_id = file_info['id']
            file_name = file_info['name']
            mime_type = file_info['mimeType']

            logger.info(f"Loading file: {file_name}")

            # ファイルの内容をダウンロード
            content = self.download_file_content(file_id, mime_type)

            if not content:
                logger.warning(f"Empty content for file: {file_name}")
                return None

            # Documentオブジェクトを作成
            document = Document(
                page_content=content,
                metadata={
                    "source": "google_drive",
                    "file_id": file_id,
                    "file_name": file_name,
                    "mime_type": mime_type,
                    "modified_time": file_info.get('modifiedTime', ''),
                    "web_view_link": file_info.get('webViewLink', ''),
                    "title": file_name
                }
            )

            logger.info(f"Loaded {len(content)} characters from {file_name}")
            return document

        except Exception as e:
            logger.error(f"Failed to load file {file_info.get('name', 'unknown')}: {e}")
            return None

    def load_folder(self, folder_id: str = None) -> list:
        """
        フォルダ内の全ファイルを読み込み

        Args:
            folder_id: Google DriveのフォルダID（Noneの場合は全ファイル）

        Returns:
            list: Documentオブジェクトのリスト
        """
        files = self.list_files_in_folder(folder_id)

        documents = []
        for file_info in files:
            doc = self.load_file(file_info)
            if doc:
                documents.append(doc)

        logger.info(f"Total: Loaded {len(documents)} documents from Google Drive")
        return documents

    def _is_garbled(self, text: str) -> bool:
        """文字化けテキストかどうかを判定（印字可能な文字の比率で判断）"""
        if not text:
            return True
        printable = sum(1 for c in text if c.isprintable() or c in '\n\t ')
        ratio = printable / len(text)
        return ratio < 0.7

    def _clean_text(self, text: str) -> str:
        """不正なUnicode文字を除去"""
        import re
        # サロゲートペア（U+D800〜U+DFFF）を除去
        cleaned = re.sub(r'[\ud800-\udfff]', '', text)
        # 制御文字を除去（改行・タブは除く）
        cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', cleaned)
        return cleaned

    def save_to_pinecone(self, documents: list) -> bool:
        """
        ドキュメントをPineconeに保存。文字化けチェック後に共通ストレージへ委譲。

        Args:
            documents: Documentオブジェクトのリスト

        Returns:
            bool: 成功したかどうか
        """
        cleaned_docs = []
        for doc in documents:
            cleaned_content = self._clean_text(doc.page_content)
            if self._is_garbled(cleaned_content):
                logger.warning(f"Skipping garbled document: {doc.metadata.get('file_name', 'unknown')}")
                continue
            cleaned_docs.append(Document(page_content=cleaned_content, metadata=doc.metadata))

        if not cleaned_docs:
            logger.warning("No valid documents after garble check")
            return False

        return _save_to_pinecone(cleaned_docs)


def load_documents_from_google_drive(folder_id: str = None):
    """
    Google Driveからドキュメントを読み込んでPineconeに保存

    Args:
        folder_id: Google DriveのフォルダID

    Returns:
        tuple: (成功したかどうか, ドキュメント数)
    """
    loader = GoogleDriveLoader()
    documents = loader.load_folder(folder_id)

    if documents:
        success = loader.save_to_pinecone(documents)
        return success, len(documents)
    return False, 0


if __name__ == "__main__":
    import logging
    logging.basicConfig(level=logging.INFO)

    folder_id = settings.google_drive_folder_id if hasattr(settings, 'google_drive_folder_id') else None

    print(f"\nGoogle Driveフォルダ: {folder_id or 'すべてのファイル'}")

    success, count = load_documents_from_google_drive(folder_id)

    if success:
        print(f"\n完了: {count} 件のドキュメントをPineconeに登録しました")
    else:
        print("\nエラー: ドキュメントの登録に失敗しました")
